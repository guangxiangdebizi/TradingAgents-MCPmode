#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
北方华创分析报告生成工具
专门处理新版JSON格式的agents_progress结构
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，无法生成Word文档")
    print("请运行: pip install python-docx")


def load_json_data(file_path: str) -> Dict[str, Any]:
    """加载JSON数据文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"错误: 文件不存在 - {file_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"错误: JSON格式错误 - {e}")
        sys.exit(1)
    except Exception as e:
        print(f"错误: 读取文件失败 - {e}")
        sys.exit(1)


def extract_agent_results(agents_progress: Dict[str, Any]) -> Dict[str, str]:
    """从agents_progress中提取各智能体的完整结果"""
    agent_results = {}
    
    # 定义智能体名称映射
    agent_names = {
        'market_analyst': '📊 市场技术分析师',
        'sentiment_analyst': '💭 市场情绪分析师', 
        'news_analyst': '📰 新闻事件分析师',
        'fundamentals_analyst': '🏢 基本面分析师',
        'bull_researcher': '📈 看涨研究员',
        'bear_researcher': '📉 看跌研究员',
        'research_manager': '👔 研究经理',
        'trader': '💼 交易员',
        'aggressive_risk_analyst': '⚡ 激进风险分析师',
        'safe_risk_analyst': '🛡️ 保守风险分析师',
        'risk_manager': '⚖️ 风险管理经理'
    }
    
    for agent_key, agent_data in agents_progress.items():
        if agent_key in agent_names and 'results' in agent_data:
            agent_name = agent_names[agent_key]
            
            # 合并所有结果
            all_results = []
            for result in agent_data['results']:
                if 'result' in result and result.get('success', True):
                    all_results.append(result['result'])
            
            if all_results:
                agent_results[agent_name] = '\n\n'.join(all_results)
            elif agent_data.get('status') == 'failed':
                # 处理失败的情况
                error_msg = "分析失败"
                if agent_data.get('results'):
                    for result in agent_data['results']:
                        if not result.get('success', True) and 'result' in result:
                            error_msg = f"分析失败: {result['result']}"
                            break
                agent_results[agent_name] = error_msg
    
    return agent_results


def generate_markdown_content(data: Dict[str, Any]) -> str:
    """生成完整的Markdown内容"""
    content = []
    
    # 标题和基本信息
    session_id = data.get('session_id', '未知会话')
    content.append(f"# {session_id} - 完整分析报告")
    content.append(f"\n**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    content.append(f"**会话ID**: {session_id}")
    content.append(f"**分析状态**: {data.get('status', '未知')}")
    
    if 'start_time' in data:
        content.append(f"**开始时间**: {data['start_time']}")
    
    content.append("\n---\n")
    
    # 用户查询
    if 'user_query' in data and data['user_query']:
        content.append("## 📋 用户查询")
        content.append(f"**问题**: {data['user_query']}\n")
    
    # 处理agents_progress
    if 'agents_progress' in data:
        agent_results = extract_agent_results(data['agents_progress'])
        
        for agent_name, result_content in agent_results.items():
            content.append(f"## {agent_name}")
            content.append("")
            content.append(result_content)
            content.append("\n---\n")
    
    # 错误信息
    if 'errors' in data and data['errors']:
        content.append("## ❌ 错误信息")
        for error in data['errors']:
            content.append(f"- {error}")
        content.append("")
    
    return "\n".join(content)


def save_markdown(content: str, output_path: str) -> str:
    """保存Markdown文件"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"✅ Markdown报告已保存到: {output_path}")
    return output_path


def save_docx(markdown_content: str, output_path: str, title: str) -> str:
    """保存Word文档，改进版本以处理长内容"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx未安装，无法生成Word文档")
    
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加标题
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加生成时间
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # 空行
    
    # 解析markdown内容并添加到文档
    lines = markdown_content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Inches(0.1)
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        line = line.strip()
        
        if not line:
            doc.add_paragraph()  # 添加空行
            continue
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # 处理粗体文本
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        # 处理列表项
        elif line.startswith('- ') or line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        # 处理分隔线
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        # 处理表格标记（简单处理）
        elif '|' in line and line.count('|') >= 2:
            # 简单的表格行处理
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                p = doc.add_paragraph()
                p.add_run(' | '.join(cells))
        # 普通段落
        else:
            # 分割长段落
            if len(line) > 1000:
                # 按句号分割长段落
                sentences = line.split('。')
                for i, sentence in enumerate(sentences):
                    if sentence.strip():
                        if i < len(sentences) - 1:
                            doc.add_paragraph(sentence.strip() + '。')
                        else:
                            doc.add_paragraph(sentence.strip())
            else:
                doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"✅ Word文档已保存到: {output_path}")
    return output_path


def main():
    # 输入文件路径
    input_file = r"c:\Users\26214\Desktop\MyProject\TradingAgents-MCPmode\progress_logs\session_分析一下北方华创咋样.json"
    
    # 检查文件是否存在
    if not Path(input_file).exists():
        print(f"错误: 文件不存在 - {input_file}")
        sys.exit(1)
    
    # 加载数据
    print(f"正在加载数据文件: {input_file}")
    data = load_json_data(input_file)
    
    # 生成报告标题
    session_id = data.get('session_id', '北方华创分析')
    title = f"{session_id} - 完整分析报告"
    
    # 生成Markdown内容
    print("正在生成Markdown内容...")
    markdown_content = generate_markdown_content(data)
    
    # 生成文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"北方华创完整分析报告_{timestamp}"
    
    # 保存Markdown文件
    md_path = f"{base_name}.md"
    save_markdown(markdown_content, md_path)
    
    # 保存Word文档
    if DOCX_AVAILABLE:
        print("正在生成Word文档...")
        docx_path = f"{base_name}.docx"
        save_docx(markdown_content, docx_path, title)
        
        # 显示文件信息
        docx_file = Path(docx_path)
        if docx_file.exists():
            file_size = docx_file.stat().st_size
            if file_size < 1024:
                size_str = f"{file_size} B"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / (1024 * 1024):.1f} MB"
            print(f"📦 Word文档大小: {size_str}")
    else:
        print("⚠️ 无法生成Word文档，请安装python-docx: pip install python-docx")
    
    print("\n🎉 报告生成完成!")
    print(f"📄 Markdown文件: {Path(md_path).absolute()}")
    if DOCX_AVAILABLE:
        print(f"📄 Word文档: {Path(docx_path).absolute()}")


if __name__ == '__main__':
    main()