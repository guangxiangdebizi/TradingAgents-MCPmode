from typing import Dict, Any, Optional
import json
import markdown
from datetime import datetime
import os
from pathlib import Path


try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX export will not be available.")

try:
    import pdfkit
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfkit not installed. PDF export will not be available.")


class ReportGenerator:
    """报告生成器 - 支持将JSON数据转换为多种格式的报告"""
    
    def __init__(self):
        self.supported_formats = ['markdown', 'md', 'html']
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['docx', 'word'])
        if PDF_AVAILABLE:
            self.supported_formats.extend(['pdf'])
    
    def generate_report(self, data: Dict[str, Any], output_format: str = 'markdown', 
                       output_path: Optional[str] = None, title: str = "交易分析报告") -> str:
        """生成报告
        
        Args:
            data: 要转换的JSON数据
            output_format: 输出格式 ('markdown', 'docx', 'pdf')
            output_path: 输出文件路径，如果为None则自动生成
            title: 报告标题
            
        Returns:
            生成的文件路径
        """
        if output_format.lower() not in self.supported_formats:
            raise ValueError(f"不支持的格式: {output_format}. 支持的格式: {self.supported_formats}")
        
        # 生成markdown内容
        markdown_content = self._generate_markdown_content(data, title)
        
        # 确定输出路径
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"trading_report_{timestamp}.{output_format.lower()}"
        
        # 根据格式生成文件
        if output_format.lower() in ['markdown', 'md']:
            return self._save_markdown(markdown_content, output_path)
        elif output_format.lower() == 'html':
            return self._save_html(markdown_content, output_path, title)
        elif output_format.lower() in ['docx', 'word'] and DOCX_AVAILABLE:
            return self._save_docx(markdown_content, output_path, title)
        elif output_format.lower() == 'pdf' and PDF_AVAILABLE:
            return self._save_pdf(markdown_content, output_path, title)
        else:
            raise RuntimeError(f"格式 {output_format} 不可用，请检查相关依赖是否已安装")
    
    def _generate_markdown_content(self, data: Dict[str, Any], title: str) -> str:
        """生成markdown内容"""
        content = []
        content.append(f"# {title}")
        content.append(f"\n**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        # 用户查询
        if 'user_query' in data:
            content.append("## 📋 用户查询")
            content.append(f"**问题**: {data['user_query']}\n")
        
        # 市场分析报告
        if 'market_report' in data and data['market_report']:
            content.append("## 📊 市场技术分析")
            content.append(self._format_report_content(data['market_report']))
            content.append("")
        
        # 情绪分析报告
        if 'sentiment_report' in data and data['sentiment_report']:
            content.append("## 💭 市场情绪分析")
            content.append(self._format_report_content(data['sentiment_report']))
            content.append("")
        
        # 新闻分析报告
        if 'news_report' in data and data['news_report']:
            content.append("## 📰 新闻事件分析")
            content.append(self._format_report_content(data['news_report']))
            content.append("")
        
        # 基本面分析报告
        if 'fundamentals_report' in data and data['fundamentals_report']:
            content.append("## 🏢 基本面分析")
            content.append(self._format_report_content(data['fundamentals_report']))
            content.append("")
        
        # 投资辩论历史
        if 'investment_debate_state' in data and data['investment_debate_state'].get('history'):
            content.append("## 🎯 投资策略辩论")
            debate_history = data['investment_debate_state']['history']
            content.append(self._format_debate_content(debate_history))
            content.append("")
        
        # 研究经理决策
        if 'research_decision' in data and data['research_decision']:
            content.append("## 👔 投资决策")
            content.append(self._format_report_content(data['research_decision']))
            content.append("")
        
        # 交易执行计划
        if 'execution_plan' in data and data['execution_plan']:
            content.append("## 📈 交易执行计划")
            content.append(self._format_report_content(data['execution_plan']))
            content.append("")
        
        # 风险管理辩论
        if 'risk_debate_state' in data and data['risk_debate_state'].get('history'):
            content.append("## ⚠️ 风险管理分析")
            risk_history = data['risk_debate_state']['history']
            content.append(self._format_debate_content(risk_history))
            content.append("")
        
        # 最终风险决策
        if 'final_risk_decision' in data and data['final_risk_decision']:
            content.append("## 🛡️ 最终风险决策")
            content.append(self._format_report_content(data['final_risk_decision']))
            content.append("")
        
        # 错误信息
        if 'errors' in data and data['errors']:
            content.append("## ❌ 错误信息")
            for error in data['errors']:
                content.append(f"- {error}")
            content.append("")
        
        # 消息历史
        if 'messages' in data and data['messages']:
            content.append("## 📝 处理日志")
            for i, message in enumerate(data['messages'], 1):
                content.append(f"### 步骤 {i}")
                content.append(f"```\n{message}\n```")
            content.append("")
        
        return "\n".join(content)
    
    def _format_report_content(self, content: str) -> str:
        """格式化报告内容"""
        if not content:
            return "暂无数据"
        
        # 简单的格式化处理
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是标题格式
            if line.endswith('：') or line.endswith(':'):
                formatted_lines.append(f"**{line}**")
            elif line.startswith('- ') or line.startswith('• '):
                formatted_lines.append(line)
            elif line.startswith(('1.', '2.', '3.', '4.', '5.')):
                formatted_lines.append(line)
            else:
                formatted_lines.append(line)
        
        return "\n".join(formatted_lines)
    
    def _format_debate_content(self, debate_history: str) -> str:
        """格式化辩论内容"""
        if not debate_history:
            return "暂无辩论记录"
        
        # 分割不同发言者的内容
        sections = debate_history.split('【')
        formatted_sections = []
        
        for section in sections:
            if not section.strip():
                continue
            
            if '】' in section:
                speaker, content = section.split('】', 1)
                speaker = speaker.strip()
                content = content.strip()
                
                formatted_sections.append(f"### {speaker}")
                formatted_sections.append(content)
                formatted_sections.append("")
            else:
                formatted_sections.append(section.strip())
        
        return "\n".join(formatted_sections)
    
    def _save_markdown(self, content: str, output_path: str) -> str:
        """保存markdown文件"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"📄 Markdown报告已保存到: {output_path}")
        return output_path
    
    def _save_html(self, markdown_content: str, output_path: str, title: str) -> str:
        """保存HTML文件"""
        # 将markdown转换为HTML
        html_body = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
        
        # 创建完整的HTML文档
        html_template = self._get_html_template(title)
        full_html = html_template.replace('{{{{CONTENT}}}}', html_body)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        print(f"🌐 HTML报告已保存到: {output_path}")
        return output_path
    
    def _get_html_template(self, title: str) -> str:
        """获取HTML模板"""
        return f'''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', 'Helvetica Neue', Helvetica, Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            overflow: hidden;
        }}
        
        .header {{
            background: linear-gradient(135deg, #2c3e50 0%, #3498db 100%);
            color: white;
            padding: 40px;
            text-align: center;
        }}
        
        .header h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
            font-weight: 300;
        }}
        
        .header .timestamp {{
            opacity: 0.8;
            font-size: 1.1em;
        }}
        
        .content {{
            padding: 40px;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 30px;
            margin-bottom: 15px;
            color: #2c3e50;
        }}
        
        h1 {{
            font-size: 2.2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        
        h2 {{
            font-size: 1.8em;
            color: #e74c3c;
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        
        h2:before {{
            content: "";
            width: 4px;
            height: 30px;
            background: #e74c3c;
            border-radius: 2px;
        }}
        
        h3 {{
            font-size: 1.4em;
            color: #8e44ad;
        }}
        
        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}
        
        strong {{
            color: #2c3e50;
            font-weight: 600;
        }}
        
        ul, ol {{
            margin-left: 30px;
            margin-bottom: 15px;
        }}
        
        li {{
            margin-bottom: 8px;
        }}
        
        code {{
            background: #f8f9fa;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
            color: #e74c3c;
        }}
        
        pre {{
            background: #2c3e50;
            color: #ecf0f1;
            padding: 20px;
            border-radius: 8px;
            overflow-x: auto;
            margin: 20px 0;
            border-left: 4px solid #3498db;
        }}
        
        pre code {{
            background: none;
            color: inherit;
            padding: 0;
        }}
        
        blockquote {{
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin: 20px 0;
            background: #f8f9fa;
            padding: 15px 20px;
            border-radius: 0 8px 8px 0;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: white;
            border-radius: 8px;
            overflow: hidden;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        
        th, td {{
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        
        th {{
            background: #3498db;
            color: white;
            font-weight: 600;
        }}
        
        tr:hover {{
            background: #f5f5f5;
        }}
        
        .section {{
            margin: 30px 0;
            padding: 25px;
            background: #f8f9fa;
            border-radius: 10px;
            border-left: 5px solid #3498db;
        }}
        
        .error-section {{
            border-left-color: #e74c3c;
            background: #fdf2f2;
        }}
        
        .warning-section {{
            border-left-color: #f39c12;
            background: #fef9e7;
        }}
        
        .success-section {{
            border-left-color: #27ae60;
            background: #f0f9f4;
        }}
        
        .footer {{
            background: #34495e;
            color: white;
            text-align: center;
            padding: 20px;
            font-size: 0.9em;
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 10px;
            }}
            
            .header {{
                padding: 20px;
            }}
            
            .header h1 {{
                font-size: 2em;
            }}
            
            .content {{
                padding: 20px;
            }}
            
            h2 {{
                font-size: 1.5em;
            }}
        }}
        
        .highlight {{
            background: linear-gradient(120deg, #a8edea 0%, #fed6e3 100%);
            padding: 2px 6px;
            border-radius: 4px;
            font-weight: 500;
        }}
        
        .badge {{
            display: inline-block;
            padding: 4px 8px;
            background: #3498db;
            color: white;
            border-radius: 12px;
            font-size: 0.8em;
            margin-right: 5px;
        }}
        
        .badge.success {{ background: #27ae60; }}
        .badge.warning {{ background: #f39c12; }}
        .badge.error {{ background: #e74c3c; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{title}</h1>
            <div class="timestamp">生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</div>
        </div>
        <div class="content">
            {{{{CONTENT}}}}
        </div>
        <div class="footer">
            <p>📊 TradingAgents-MCPmode 智能交易分析系统 | 🤖 多智能体协作分析报告</p>
        </div>
    </div>
</body>
</html>
        '''
    
    def _save_docx(self, markdown_content: str, output_path: str, title: str) -> str:
        """保存DOCX文件"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx未安装，无法生成DOCX文件")
        
        doc = Document()
        
        # 添加标题
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()  # 空行
        
        # 解析markdown内容并添加到文档
        lines = markdown_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_paragraph is not None:
                    doc.add_paragraph()  # 添加空行
                continue
            
            # 处理标题
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            # 处理粗体文本
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            # 处理列表项
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # 处理代码块
            elif line.startswith('```'):
                continue  # 跳过代码块标记
            # 普通段落
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
        logger.info(f"DOCX报告已保存到: {output_path}")
        return output_path
    
    def _save_pdf(self, markdown_content: str, output_path: str, title: str) -> str:
        """保存PDF文件"""
        if not PDF_AVAILABLE:
            raise RuntimeError("pdfkit未安装，无法生成PDF文件")
        
        # 将markdown转换为HTML
        html_content = markdown.markdown(markdown_content)
        
        # 添加CSS样式
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{
                    font-family: 'Microsoft YaHei', Arial, sans-serif;
                    line-height: 1.6;
                    margin: 40px;
                    color: #333;
                }}
                h1, h2, h3 {{
                    color: #2c3e50;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }}
                h1 {{ font-size: 28px; }}
                h2 {{ font-size: 24px; }}
                h3 {{ font-size: 20px; }}
                p {{ margin-bottom: 15px; }}
                ul, ol {{ margin-left: 20px; }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 4px;
                    border-radius: 3px;
                }}
                pre {{
                    background-color: #f4f4f4;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # 配置PDF选项
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }
        
        # 生成PDF
        pdfkit.from_string(styled_html, output_path, options=options)
        logger.info(f"PDF报告已保存到: {output_path}")
        return output_path
    
    def get_supported_formats(self) -> list:
        """获取支持的输出格式"""
        return self.supported_formats.copy()
    
    def check_dependencies(self) -> Dict[str, bool]:
        """检查依赖项是否可用"""
        return {
            'markdown': True,  # 内置支持
            'docx': DOCX_AVAILABLE,
            'pdf': PDF_AVAILABLE
        }


def generate_trading_report(json_data: Dict[str, Any], output_format: str = 'markdown', 
                          output_path: Optional[str] = None, title: str = "交易分析报告") -> str:
    """便捷函数：生成交易分析报告
    
    Args:
        json_data: 交易分析的JSON数据
        output_format: 输出格式 ('markdown', 'docx', 'pdf')
        output_path: 输出文件路径
        title: 报告标题
        
    Returns:
        生成的文件路径
    """
    generator = ReportGenerator()
    return generator.generate_report(json_data, output_format, output_path, title)


if __name__ == "__main__":
    # 示例用法
    sample_data = {
        "user_query": "分析苹果公司股票",
        "market_report": "技术分析显示股价处于上升趋势...",
        "sentiment_report": "市场情绪整体偏向乐观...",
        "investment_debate_state": {
            "history": "【看涨研究员 第1轮】:\n基于技术分析，建议买入...\n\n【看跌研究员 第1轮】:\n存在以下风险因素..."
        }
    }
    
    generator = ReportGenerator()
    print(f"支持的格式: {generator.get_supported_formats()}")
    print(f"依赖项状态: {generator.check_dependencies()}")
    
    # 生成markdown报告
    try:
        output_file = generator.generate_report(sample_data, 'markdown', 'sample_report.md')
        print(f"报告已生成: {output_file}")
    except Exception as e:
        print(f"生成报告失败: {e}")