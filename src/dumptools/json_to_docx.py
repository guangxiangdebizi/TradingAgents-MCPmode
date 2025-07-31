#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JSON to DOCX Converter
专门为dump文件夹下的JSON文档导出为DOCX的工具
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️ python-docx未安装，DOCX导出功能不可用")
    print("📦 请安装: pip install python-docx")


class JSONToDocxConverter:
    """JSON转DOCX转换器"""
    
    def __init__(self, dump_dir: str = "src/dump"):
        """初始化转换器
        
        Args:
            dump_dir: dump文件夹路径
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，请先安装相关依赖")
            
        self.dump_dir = Path(dump_dir)
        self.output_dir = Path("docx_reports")
        
        # 确保输出目录存在
        self.output_dir.mkdir(exist_ok=True)
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        try:
            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Microsoft YaHei'
            font.size = Pt(11)
            
            # 创建标题样式
            try:
                title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = 'Microsoft YaHei'
                title_font.size = Pt(20)
                title_font.bold = True
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(12)
            except:
                pass  # 样式可能已存在
            
            # 创建章节标题样式
            try:
                heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                heading1_font = heading1_style.font
                heading1_font.name = 'Microsoft YaHei'
                heading1_font.size = Pt(16)
                heading1_font.bold = True
                heading1_style.paragraph_format.space_after = Pt(12)
            except:
                pass
            
            # 创建子标题样式
            try:
                heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
                heading2_font = heading2_style.font
                heading2_font.name = 'Microsoft YaHei'
                heading2_font.size = Pt(14)
                heading2_font.bold = True
                heading2_style.paragraph_format.space_after = Pt(10)
            except:
                pass
            
            # 创建代码样式
            try:
                code_style = doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
                code_font = code_style.font
                code_font.name = 'Consolas'
                code_font.size = Pt(9)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.space_after = Pt(6)
            except:
                pass
                
        except Exception as e:
            print(f"⚠️ 设置文档样式时出错: {e}")
    
    def _add_emoji_support(self, paragraph, text: str):
        """添加表情符号支持"""
        try:
            run = paragraph.add_run(text)
            # 设置字体以支持表情符号
            run.font.name = 'Segoe UI Emoji'
            return run
        except:
            # 如果失败，使用默认字体
            return paragraph.add_run(text)
    
    def convert_json_to_docx(self, json_file_path: str) -> Optional[str]:
        """将JSON文件转换为DOCX
        
        Args:
            json_file_path: JSON文件路径
            
        Returns:
            生成的DOCX文件路径，失败返回None
        """
        try:
            # 读取JSON文件
            with open(json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 创建文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_styles(doc)
            
            # 生成DOCX内容
            self._generate_docx_content(doc, data)
            
            # 生成输出文件名
            json_filename = Path(json_file_path).stem
            output_file = self.output_dir / f"{json_filename}.docx"
            
            # 保存文档
            doc.save(str(output_file))
            
            print(f"✅ DOCX报告已生成: {output_file}")
            return str(output_file)
            
        except Exception as e:
            print(f"❌ DOCX转换失败: {e}")
            return None
    
    def _generate_docx_content(self, doc: Document, data: Dict[str, Any]):
        """生成DOCX内容"""
        # 标题
        session_id = data.get('session_id', 'Unknown')
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(f"交易分析报告 - {session_id}")
        title_run.font.name = 'Microsoft YaHei'
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
        
        # 基本信息
        heading = doc.add_paragraph()
        emoji_run = heading.add_run("📋 ")
        emoji_run.font.name = 'Segoe UI Emoji'
        text_run = heading.add_run("基本信息")
        text_run.font.name = 'Microsoft YaHei'
        text_run.font.size = Pt(16)
        text_run.bold = True
        
        basic_info = [
            f"会话ID: {data.get('session_id', 'N/A')}",
            f"创建时间: {data.get('created_at', 'N/A')}",
            f"更新时间: {data.get('updated_at', 'N/A')}",
            f"状态: {data.get('status', 'N/A')}"
        ]
        
        for info in basic_info:
            p = doc.add_paragraph()
            bullet_run = p.add_run("• ")
            bullet_run.font.name = 'Microsoft YaHei'
            info_run = p.add_run(info)
            info_run.font.name = 'Microsoft YaHei'
        
        doc.add_paragraph()  # 空行
        
        # 用户查询
        if 'user_query' in data and data['user_query']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("🔍 ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("用户查询")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            query_p = doc.add_paragraph()
            quote_run1 = query_p.add_run("❝ ")
            quote_run1.font.name = 'Segoe UI Emoji'
            query_run = query_p.add_run(data['user_query'])
            query_run.font.name = 'Microsoft YaHei'
            quote_run2 = query_p.add_run(" ❞")
            quote_run2.font.name = 'Segoe UI Emoji'
            
            doc.add_paragraph()  # 空行
        
        # 智能体执行情况
        if 'agents' in data and data['agents']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("🤖 ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("智能体执行情况")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for agent in data['agents']:
                agent_name = agent.get('agent_name', 'Unknown Agent')
                
                # 智能体名称
                agent_heading = doc.add_paragraph()
                agent_run = agent_heading.add_run(agent_name)
                agent_run.font.name = 'Microsoft YaHei'
                agent_run.font.size = Pt(14)
                agent_run.bold = True
                
                # 智能体信息
                agent_info = [
                    f"状态: {agent.get('status', 'N/A')}",
                    f"开始时间: {agent.get('start_time', 'N/A')}",
                ]
                
                if agent.get('end_time'):
                    agent_info.append(f"结束时间: {agent.get('end_time')}")
                
                agent_info.append(f"执行结果: {agent.get('result', 'N/A')}")
                
                for info in agent_info:
                    p = doc.add_paragraph()
                    bullet_run = p.add_run("• ")
                    bullet_run.font.name = 'Microsoft YaHei'
                    info_run = p.add_run(info)
                    info_run.font.name = 'Microsoft YaHei'
                
                # 执行内容
                if agent.get('action'):
                    action_heading = doc.add_paragraph()
                    action_title = action_heading.add_run("执行内容:")
                    action_title.font.name = 'Microsoft YaHei'
                    action_title.bold = True
                    
                    # 将长文本分段处理
                    action_text = str(agent['action'])
                    if len(action_text) > 1000:
                        action_text = action_text[:1000] + "..."
                    
                    action_p = doc.add_paragraph()
                    action_run = action_p.add_run(action_text)
                    action_run.font.name = 'Consolas'
                    action_run.font.size = Pt(9)
                    action_p.paragraph_format.left_indent = Inches(0.5)
        
        # 阶段信息
        if 'stages' in data and data['stages']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("📊 ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("执行阶段")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for i, stage in enumerate(data['stages'], 1):
                stage_heading = doc.add_paragraph()
                stage_run = stage_heading.add_run(f"阶段 {i}")
                stage_run.font.name = 'Microsoft YaHei'
                stage_run.font.size = Pt(14)
                stage_run.bold = True
                
                stage_p = doc.add_paragraph()
                content_run = stage_p.add_run(f"内容: {stage}")
                content_run.font.name = 'Microsoft YaHei'
        
        # MCP调用情况
        if 'mcp_calls' in data and data['mcp_calls']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("🔧 ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("MCP工具调用")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for i, call in enumerate(data['mcp_calls'], 1):
                call_heading = doc.add_paragraph()
                call_run = call_heading.add_run(f"调用 {i}")
                call_run.font.name = 'Microsoft YaHei'
                call_run.font.size = Pt(14)
                call_run.bold = True
                
                call_info = [
                    f"工具: {call.get('tool', 'N/A')}",
                    f"时间: {call.get('timestamp', 'N/A')}"
                ]
                if call.get('result'):
                    call_info.append(f"结果: {call['result']}")
                
                for info in call_info:
                    p = doc.add_paragraph()
                    bullet_run = p.add_run("• ")
                    bullet_run.font.name = 'Microsoft YaHei'
                    info_run = p.add_run(info)
                    info_run.font.name = 'Microsoft YaHei'
        
        # 错误信息
        if 'errors' in data and data['errors']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("❌ ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("错误信息")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for error in data['errors']:
                p = doc.add_paragraph()
                bullet_run = p.add_run("• ")
                bullet_run.font.name = 'Microsoft YaHei'
                error_run = p.add_run(error)
                error_run.font.name = 'Microsoft YaHei'
        
        # 警告信息
        if 'warnings' in data and data['warnings']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("⚠️ ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("警告信息")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for warning in data['warnings']:
                p = doc.add_paragraph()
                bullet_run = p.add_run("• ")
                bullet_run.font.name = 'Microsoft YaHei'
                warning_run = p.add_run(warning)
                warning_run.font.name = 'Microsoft YaHei'
        
        # 最终结果
        if 'final_results' in data and data['final_results']:
            heading = doc.add_paragraph()
            emoji_run = heading.add_run("🎯 ")
            emoji_run.font.name = 'Segoe UI Emoji'
            text_run = heading.add_run("最终结果")
            text_run.font.name = 'Microsoft YaHei'
            text_run.font.size = Pt(16)
            text_run.bold = True
            
            for key, value in data['final_results'].items():
                result_heading = doc.add_paragraph()
                result_run = result_heading.add_run(f"{key}")
                result_run.font.name = 'Microsoft YaHei'
                result_run.font.size = Pt(14)
                result_run.bold = True
                
                result_text = str(value)
                if len(result_text) > 1000:
                    result_text = result_text[:1000] + "..."
                
                result_p = doc.add_paragraph()
                result_content = result_p.add_run(result_text)
                result_content.font.name = 'Consolas'
                result_content.font.size = Pt(9)
                result_p.paragraph_format.left_indent = Inches(0.5)
        
        # 分隔线和生成时间戳
        doc.add_paragraph()
        separator_p = doc.add_paragraph()
        separator_run = separator_p.add_run("─" * 50)
        separator_run.font.name = 'Microsoft YaHei'
        
        timestamp_p = doc.add_paragraph()
        timestamp_run = timestamp_p.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_run.font.name = 'Microsoft YaHei'
        timestamp_run.italic = True
    
    def convert_latest_json(self) -> Optional[str]:
        """转换最新的JSON文件
        
        Returns:
            生成的DOCX文件路径，失败返回None
        """
        try:
            # 查找dump目录下的所有JSON文件
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"❌ 在 {self.dump_dir} 目录下未找到JSON文件")
                return None
            
            # 找到最新的文件
            latest_json = max(json_files, key=lambda f: f.stat().st_mtime)
            print(f"📄 找到最新的JSON文件: {latest_json.name}")
            
            # 转换为DOCX
            return self.convert_json_to_docx(str(latest_json))
            
        except Exception as e:
            print(f"❌ DOCX转换过程中发生错误: {e}")
            return None
    
    def convert_all_json(self) -> List[str]:
        """转换所有JSON文件
        
        Returns:
            生成的DOCX文件路径列表
        """
        try:
            # 查找dump目录下的所有JSON文件
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"❌ 在 {self.dump_dir} 目录下未找到JSON文件")
                return []
            
            results = []
            for json_file in json_files:
                print(f"📄 转换文件: {json_file.name}")
                result = self.convert_json_to_docx(str(json_file))
                if result:
                    results.append(result)
            
            return results
            
        except Exception as e:
            print(f"❌ DOCX批量转换过程中发生错误: {e}")
            return []


def main():
    """主函数 - 命令行工具"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="JSON to DOCX Converter - 将dump文件夹下的JSON文档转换为DOCX"
    )
    parser.add_argument("-f", "--file", help="指定要转换的JSON文件路径")
    parser.add_argument("-l", "--latest", action="store_true", help="转换最新的JSON文件")
    parser.add_argument("-a", "--all", action="store_true", help="转换所有JSON文件")
    parser.add_argument("-d", "--dump-dir", default="src/dump", help="dump文件夹路径")
    
    args = parser.parse_args()
    
    try:
        converter = JSONToDocxConverter(args.dump_dir)
        
        if args.all:
            # 转换所有文件
            results = converter.convert_all_json()
            if results:
                print(f"🎉 批量转换完成，共生成 {len(results)} 个DOCX文件")
            else:
                print("❌ 批量转换失败")
        
        elif args.latest:
            # 转换最新文件
            result = converter.convert_latest_json()
            if result:
                print(f"🎉 转换完成: {result}")
        
        elif args.file:
            # 转换指定文件
            if os.path.exists(args.file):
                result = converter.convert_json_to_docx(args.file)
                if result:
                    print(f"🎉 转换完成: {result}")
            else:
                print(f"❌ 文件不存在: {args.file}")
        
        else:
            # 默认转换最新文件
            result = converter.convert_latest_json()
            if result:
                print(f"🎉 转换完成: {result}")
                
    except ImportError as e:
        print(f"❌ 依赖缺失: {e}")
        print("📦 请安装必要的依赖:")
        print("   pip install python-docx")


if __name__ == "__main__":
    main()
