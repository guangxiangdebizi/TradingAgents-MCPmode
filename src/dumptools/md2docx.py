#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter
åŸºäºjson_to_markdown.pyçš„ä¸¤æ­¥è½¬æ¢ï¼šJSON â†’ Markdown â†’ DOCX
"""

import os
import sys
from pathlib import Path
from typing import Optional, List
import argparse
from datetime import datetime
import re

# å¯¼å…¥markdownè½¬æ¢å™¨
from json_to_markdown import JSONToMarkdownConverter

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("âŒ ç¼ºå°‘ä¾èµ–åŒ…ï¼Œè¯·å®‰è£…ï¼špip install python-docx")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Markdownè½¬DOCXè½¬æ¢å™¨"""
    
    def __init__(self, dump_dir: str = "src/dump"):
        """åˆå§‹åŒ–è½¬æ¢å™¨
        
        Args:
            dump_dir: dumpæ–‡ä»¶å¤¹è·¯å¾„
        """
        self.dump_dir = Path(dump_dir)
        # ä¿æŒåŸæœ‰çš„è¾“å‡ºç›®å½•è®¾ç½®
        self.output_dir = self.dump_dir.parent / "output"
        self.output_dir.mkdir(exist_ok=True)
        
        # åˆå§‹åŒ–Markdownè½¬æ¢å™¨
        from json_to_markdown import JSONToMarkdownConverter
        self.md_converter = JSONToMarkdownConverter(str(self.dump_dir))
        
        # ç›®å½•é¡¹åˆ—è¡¨
        self.toc_entries = []
    
    def _setup_document_styles(self, doc):
        """è®¾ç½®æ–‡æ¡£æ ·å¼"""
        # è®¾ç½®é»˜è®¤å­—ä½“
        style = doc.styles['Normal']
        font = style.font
        font.name = 'å¾®è½¯é›…é»‘'
        font.size = Pt(10)
        
        # åˆ›å»ºä¸­æ–‡æ ‡é¢˜æ ·å¼
        if 'Chinese Title' not in [s.name for s in doc.styles]:
            title_style = doc.styles.add_style('Chinese Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'å¾®è½¯é›…é»‘'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)
        
        # åˆ›å»ºå„çº§æ ‡é¢˜æ ·å¼
        heading_configs = [
            ('Chinese Heading 1', 16, True),
            ('Chinese Heading 2', 14, True),
            ('Chinese Heading 3', 12, True),
            ('Chinese Heading 4', 11, True)
        ]
        
        for style_name, font_size, is_bold in heading_configs:
            if style_name not in [s.name for s in doc.styles]:
                heading_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_font = heading_style.font
                heading_font.name = 'å¾®è½¯é›…é»‘'
                heading_font.size = Pt(font_size)
                heading_font.bold = is_bold
                heading_style.paragraph_format.space_after = Pt(10)
                heading_style.paragraph_format.space_before = Pt(10)
        
        # åˆ›å»ºä»£ç æ ·å¼
        if 'Chinese Code' not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style('Chinese Code', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
        
        # åˆ›å»ºå¼•ç”¨æ ·å¼
        if 'Chinese Quote' not in [s.name for s in doc.styles]:
            quote_style = doc.styles.add_style('Chinese Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_font = quote_style.font
            quote_font.name = 'å¾®è½¯é›…é»‘'
            quote_font.size = Pt(10)
            quote_font.italic = True
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
        
        # åˆ›å»ºç›®å½•æ ·å¼
        if 'TOC Heading' not in [s.name for s in doc.styles]:
            toc_heading_style = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
            toc_heading_font = toc_heading_style.font
            toc_heading_font.name = 'å¾®è½¯é›…é»‘'
            toc_heading_font.size = Pt(16)
            toc_heading_font.bold = True
            toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_heading_style.paragraph_format.space_after = Pt(20)
    
    def _add_emoji_support(self, doc):
        """æ·»åŠ emojiå­—ä½“æ”¯æŒ"""
        try:
            # ä¸ºæ–‡æ¡£æ·»åŠ Segoe UI Emojiå­—ä½“æ”¯æŒ
            # è¿™éœ€è¦åœ¨è¿è¡Œæ—¶è®¾ç½®å­—ä½“å›é€€
            pass
        except Exception as e:
            print(f"âš ï¸ Emojiå­—ä½“è®¾ç½®å¤±è´¥: {e}")
    
    def _create_table_of_contents(self, doc):
        """åˆ›å»ºç›®å½•"""
        if not self.toc_entries:
            return
        
        # æ·»åŠ ç›®å½•æ ‡é¢˜
        toc_heading = doc.add_paragraph('ğŸ“‹ ç›®å½•', style='TOC Heading')
        
        # æ·»åŠ ç›®å½•é¡¹
        for level, title, page_num in self.toc_entries:
            toc_paragraph = doc.add_paragraph()
            
            # è®¾ç½®ç¼©è¿›
            toc_paragraph.paragraph_format.left_indent = Inches((level - 1) * 0.3)
            
            # æ·»åŠ æ ‡é¢˜æ–‡æœ¬
            run = toc_paragraph.add_run(title)
            run.font.name = 'å¾®è½¯é›…é»‘'
            
            # å°è¯•è®¾ç½®emojiå­—ä½“
            if any(ord(char) > 127 for char in title if ord(char) > 0x1F000):
                run.font.name = 'Segoe UI Emoji'
            
            # æ·»åŠ ç‚¹çº¿å’Œé¡µç ï¼ˆç®€åŒ–ç‰ˆï¼‰
            dots_run = toc_paragraph.add_run(' ' + '.' * (50 - len(title)))
            dots_run.font.name = 'å¾®è½¯é›…é»‘'
            
            page_run = toc_paragraph.add_run(f' {page_num}')
            page_run.font.name = 'å¾®è½¯é›…é»‘'
        
        # æ·»åŠ åˆ†é¡µç¬¦
        doc.add_page_break()
    
    def _parse_markdown_to_docx(self, markdown_content: str, doc):
        """è§£æMarkdownå†…å®¹å¹¶æ·»åŠ åˆ°DOCXæ–‡æ¡£"""
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # å¤„ç†æ ‡é¢˜
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading = line.lstrip('#').strip()
                
                if level == 1:
                    p = doc.add_paragraph(heading, style='Chinese Heading 1')
                elif level == 2:
                    p = doc.add_paragraph(heading, style='Chinese Heading 2')
                elif level == 3:
                    p = doc.add_paragraph(heading, style='Chinese Heading 3')
                else:
                    p = doc.add_paragraph(heading, style='Chinese Heading 4')
                
                # ä¸ºemojiè®¾ç½®ç‰¹æ®Šå­—ä½“
                for run in p.runs:
                    if any(ord(char) > 0x1F000 for char in run.text):
                        run.font.name = 'Segoe UI Emoji'
            
            # å¤„ç†å¼•ç”¨
            elif line.startswith('> '):
                quote = line[2:].strip()
                p = doc.add_paragraph(quote, style='Chinese Quote')
            
            # å¤„ç†åˆ—è¡¨
            elif line.startswith('- '):
                item = line[2:].strip()
                p = doc.add_paragraph(item, style='List Bullet')
                # è®¾ç½®å­—ä½“
                for run in p.runs:
                    run.font.name = 'å¾®è½¯é›…é»‘'
                    run.font.size = Pt(10)
            
            # å¤„ç†è¡¨æ ¼
            elif '|' in line and line.strip().startswith('|'):
                # æ”¶é›†è¡¨æ ¼è¡Œ
                table_rows = []
                table_rows.append(line.strip())
                
                # ç»§ç»­è¯»å–è¡¨æ ¼è¡Œ
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if next_line and '|' in next_line and next_line.startswith('|'):
                        table_rows.append(next_line)
                        j += 1
                    else:
                        break
                
                # è§£æè¡¨æ ¼å¹¶æ·»åŠ åˆ°æ–‡æ¡£
                if len(table_rows) > 1:
                    self._add_table_to_doc(table_rows, doc)
                    i = j - 1  # è°ƒæ•´ç´¢å¼•
            
            # å¤„ç†ä»£ç å—
            elif line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    # ä¿æŒå®Œæ•´å†…å®¹ï¼Œä¸è¿›è¡Œæˆªæ–­
                    
                    # åˆ†è¡Œæ·»åŠ ä»£ç 
                    for code_line in code_text.split('\n'):
                        p = doc.add_paragraph(code_line, style='Chinese Code')
            
            # å¤„ç†æ™®é€šæ–‡æœ¬
            else:
                if line:
                    p = doc.add_paragraph()
                    
                    # å¤„ç†ç²—ä½“æ–‡æœ¬
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # ç²—ä½“æ–‡æœ¬
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            # æ™®é€šæ–‡æœ¬
                            run = p.add_run(part)
                        
                        # è®¾ç½®å­—ä½“
                        run.font.name = 'å¾®è½¯é›…é»‘'
                        run.font.size = Pt(10)
            
            i += 1
    
    def _add_table_to_doc(self, table_rows, doc):
        """å°†è¡¨æ ¼æ·»åŠ åˆ°Wordæ–‡æ¡£"""
        # è§£æè¡¨æ ¼æ•°æ®
        table_data = []
        for row in table_rows:
            # ç§»é™¤é¦–å°¾çš„|ç¬¦å·å¹¶åˆ†å‰²
            cells = [cell.strip() for cell in row.strip('|').split('|')]
            table_data.append(cells)
        
        # è·³è¿‡åˆ†éš”è¡Œï¼ˆé€šå¸¸æ˜¯ç¬¬äºŒè¡Œï¼ŒåŒ…å«---ï¼‰
        if len(table_data) > 1 and all('---' in cell or '-' in cell for cell in table_data[1]):
            table_data.pop(1)
        
        if table_data and len(table_data) > 0:
            # åˆ›å»ºè¡¨æ ¼
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            
            if cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # å¡«å……è¡¨æ ¼æ•°æ®
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            
                            # è®¾ç½®å­—ä½“
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'å¾®è½¯é›…é»‘'
                                    run.font.size = Pt(9)
                            
                            # è¡¨å¤´åŠ ç²—
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                # æ·»åŠ æ®µè½é—´è·
                doc.add_paragraph()

    def convert_json_to_docx_via_markdown(self, json_file_path: str) -> Optional[str]:
        """é€šè¿‡Markdownä¸­é—´æ­¥éª¤å°†JSONè½¬æ¢ä¸ºDOCX
        
        Args:
            json_file_path: JSONæ–‡ä»¶è·¯å¾„
            
        Returns:
            ç”Ÿæˆçš„DOCXæ–‡ä»¶è·¯å¾„ï¼Œå¤±è´¥è¿”å›None
        """
        try:
            # ç¬¬ä¸€æ­¥ï¼šJSONè½¬Markdown
            print(f"ğŸ“„ ç¬¬ä¸€æ­¥ï¼šå°†JSONè½¬æ¢ä¸ºMarkdown...")
            md_file_path = self.md_converter.convert_json_to_markdown(json_file_path)
            
            if not md_file_path or not os.path.exists(md_file_path):
                print("âŒ Markdownè½¬æ¢å¤±è´¥")
                return None
            
            # ç¬¬äºŒæ­¥ï¼šMarkdownè½¬DOCX
            print(f"ğŸ“„ ç¬¬äºŒæ­¥ï¼šå°†Markdownè½¬æ¢ä¸ºDOCX...")
            
            # è¯»å–Markdownæ–‡ä»¶
            with open(md_file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # åˆ›å»ºæ–°çš„Wordæ–‡æ¡£
            doc = Document()
            
            # è®¾ç½®æ–‡æ¡£æ ·å¼
            self._setup_document_styles(doc)
            self._add_emoji_support(doc)
            
            # è§£æMarkdownå†…å®¹å¹¶æ·»åŠ åˆ°æ–‡æ¡£
            self._parse_markdown_to_docx(markdown_content, doc)
            
            # ç”ŸæˆDOCXæ–‡ä»¶å
            json_filename = Path(json_file_path).stem
            docx_file = self.output_dir / f"{json_filename}.docx"
            
            # ä¿å­˜æ–‡æ¡£
            doc.save(str(docx_file))
            
            print(f"âœ… DOCXæŠ¥å‘Šå·²ç”Ÿæˆ: {docx_file}")
            return str(docx_file)
            
        except Exception as e:
            print(f"âŒ è½¬æ¢å¤±è´¥: {e}")
            return None
    
    def convert_latest_json(self) -> Optional[str]:
        """è½¬æ¢æœ€æ–°çš„JSONæ–‡ä»¶
        
        Returns:
            ç”Ÿæˆçš„DOCXæ–‡ä»¶è·¯å¾„ï¼Œå¤±è´¥è¿”å›None
        """
        try:
            # æŸ¥æ‰¾dumpç›®å½•ä¸‹çš„æ‰€æœ‰JSONæ–‡ä»¶
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"âŒ åœ¨ {self.dump_dir} ç›®å½•ä¸‹æœªæ‰¾åˆ°JSONæ–‡ä»¶")
                return None
            
            # æ‰¾åˆ°æœ€æ–°çš„æ–‡ä»¶
            latest_json = max(json_files, key=lambda f: f.stat().st_mtime)
            print(f"ğŸ“„ æ‰¾åˆ°æœ€æ–°çš„JSONæ–‡ä»¶: {latest_json.name}")
            
            # è½¬æ¢ä¸ºDOCX
            return self.convert_json_to_docx_via_markdown(str(latest_json))
            
        except Exception as e:
            print(f"âŒ è½¬æ¢è¿‡ç¨‹ä¸­å‘ç”Ÿé”™è¯¯: {e}")
            return None
    
    def convert_all_json(self) -> List[str]:
        """è½¬æ¢æ‰€æœ‰JSONæ–‡ä»¶
        
        Returns:
            ç”Ÿæˆçš„DOCXæ–‡ä»¶è·¯å¾„åˆ—è¡¨
        """
        try:
            # æŸ¥æ‰¾dumpç›®å½•ä¸‹çš„æ‰€æœ‰JSONæ–‡ä»¶
            json_files = list(self.dump_dir.glob("session_*.json"))
            
            if not json_files:
                print(f"âŒ åœ¨ {self.dump_dir} ç›®å½•ä¸‹æœªæ‰¾åˆ°JSONæ–‡ä»¶")
                return []
            
            results = []
            for json_file in json_files:
                print(f"ğŸ“„ è½¬æ¢æ–‡ä»¶: {json_file.name}")
                result = self.convert_json_to_docx_via_markdown(str(json_file))
                if result:
                    results.append(result)
            
            return results
            
        except Exception as e:
            print(f"âŒ æ‰¹é‡è½¬æ¢è¿‡ç¨‹ä¸­å‘ç”Ÿé”™è¯¯: {e}")
            return []


def main():
    """ä¸»å‡½æ•° - å‘½ä»¤è¡Œå·¥å…·"""
    parser = argparse.ArgumentParser(
        description="Markdown to DOCX Converter - é€šè¿‡Markdownä¸­é—´æ­¥éª¤å°†JSONè½¬æ¢ä¸ºDOCX"
    )
    parser.add_argument("-f", "--file", help="æŒ‡å®šè¦è½¬æ¢çš„JSONæ–‡ä»¶è·¯å¾„")
    parser.add_argument("-l", "--latest", action="store_true", help="è½¬æ¢æœ€æ–°çš„JSONæ–‡ä»¶")
    parser.add_argument("-a", "--all", action="store_true", help="è½¬æ¢æ‰€æœ‰JSONæ–‡ä»¶")
    parser.add_argument("-d", "--dump-dir", default="src/dump", help="dumpæ–‡ä»¶å¤¹è·¯å¾„")
    
    args = parser.parse_args()
    
    converter = MarkdownToDocxConverter(args.dump_dir)
    
    if args.all:
        # è½¬æ¢æ‰€æœ‰æ–‡ä»¶
        results = converter.convert_all_json()
        if results:
            print(f"ğŸ‰ æ‰¹é‡è½¬æ¢å®Œæˆï¼Œå…±ç”Ÿæˆ {len(results)} ä¸ªDOCXæ–‡ä»¶")
        else:
            print("âŒ æ‰¹é‡è½¬æ¢å¤±è´¥")
    
    elif args.latest:
        # è½¬æ¢æœ€æ–°æ–‡ä»¶
        result = converter.convert_latest_json()
        if result:
            print(f"ğŸ‰ è½¬æ¢å®Œæˆ: {result}")
    
    elif args.file:
        # è½¬æ¢æŒ‡å®šæ–‡ä»¶
        if os.path.exists(args.file):
            result = converter.convert_json_to_docx_via_markdown(args.file)
            if result:
                print(f"ğŸ‰ è½¬æ¢å®Œæˆ: {result}")
        else:
            print(f"âŒ æ–‡ä»¶ä¸å­˜åœ¨: {args.file}")
    
    else:
        # é»˜è®¤è½¬æ¢æœ€æ–°æ–‡ä»¶
        result = converter.convert_latest_json()
        if result:
            print(f"ğŸ‰ è½¬æ¢å®Œæˆ: {result}")


if __name__ == "__main__":
    main()