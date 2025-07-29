#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JSON分析报告转Word文档转换器
将JSON格式的交易分析报告转换为精美的Word文档
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, url, text):
    """添加超链接到段落"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # 设置超链接样式
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def get_agent_display_info(agent_type):
    """获取智能体的显示信息"""
    agent_info = {
        'fundamental_analyst': {
            'name': '基本面分析师',
            'icon': '📊',
            'color': RGBColor(52, 152, 219)
        },
        'technical_analyst': {
            'name': '技术分析师', 
            'icon': '📈',
            'color': RGBColor(155, 89, 182)
        },
        'market_researcher': {
            'name': '市场研究员',
            'icon': '🔍',
            'color': RGBColor(46, 204, 113)
        },
        'news_researcher': {
            'name': '新闻研究员',
            'icon': '📰',
            'color': RGBColor(241, 196, 15)
        },
        'portfolio_manager': {
            'name': '投资组合经理',
            'icon': '💼',
            'color': RGBColor(230, 126, 34)
        },
        'risk_management_manager': {
            'name': '风险管理经理',
            'icon': '🛡️',
            'color': RGBColor(231, 76, 60)
        },
        'neutral_risk_analyst': {
            'name': '中性风险分析师',
            'icon': '⚖️',
            'color': RGBColor(149, 165, 166)
        },
        'senior_trader': {
            'name': '高级交易员',
            'icon': '💰',
            'color': RGBColor(26, 188, 156)
        }
    }
    return agent_info.get(agent_type, {
        'name': agent_type.replace('_', ' ').title(),
        'icon': '🤖',
        'color': RGBColor(127, 140, 141)
    })

def format_content(content):
    """格式化内容，处理特殊字符和格式"""
    if not content:
        return ""
    
    # 移除多余的空行
    lines = content.split('\n')
    formatted_lines = []
    prev_empty = False
    
    for line in lines:
        line = line.strip()
        if not line:
            if not prev_empty:
                formatted_lines.append('')
            prev_empty = True
        else:
            formatted_lines.append(line)
            prev_empty = False
    
    return '\n'.join(formatted_lines)

def extract_content_from_result(result_data):
    """从结果数据中提取实际内容"""
    if isinstance(result_data, dict):
        if 'content' in result_data:
            content = result_data['content']
            # 如果content是字符串形式的字典，尝试解析
            if isinstance(content, str) and content.startswith("{"):
                try:
                    parsed_content = eval(content)  # 注意：这里使用eval，实际应用中应该更安全
                    if isinstance(parsed_content, dict) and 'result' in parsed_content:
                        return parsed_content['result']
                except:
                    pass
            return content
        elif 'result' in result_data:
            return result_data['result']
    elif isinstance(result_data, str):
        return result_data
    return str(result_data)

def create_word_report(json_file_path):
    """创建Word格式的分析报告"""
    try:
        # 读取JSON文件
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # 添加标题
        title = doc.add_heading('📈 智能交易分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(52, 73, 94)
        title_run.font.size = Pt(24)
        
        # 添加分析目标
        if 'user_query' in data:
            target_para = doc.add_paragraph()
            target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            target_run = target_para.add_run(f"分析目标: {data['user_query']}")
            target_run.font.size = Pt(16)
            target_run.font.color.rgb = RGBColor(41, 128, 185)
            target_run.bold = True
        
        # 添加时间信息
        if 'created_at' in data:
            time_para = doc.add_paragraph()
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_run = time_para.add_run(f"生成时间: {data['created_at']}")
            time_run.font.size = Pt(12)
            time_run.font.color.rgb = RGBColor(127, 140, 141)
        
        # 添加状态信息
        if 'status' in data:
            status_para = doc.add_paragraph()
            status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            status_run = status_para.add_run(f"分析状态: {data['status']}")
            status_run.font.size = Pt(12)
            status_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # 空行
        
        # 处理智能体分析结果
        if 'agents_data' in data:
            agents_heading = doc.add_heading('🤖 智能体分析详情', level=1)
            agents_heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
            
            for agent_type, agent_data in data['agents_data'].items():
                if not agent_data:
                    continue
                
                agent_info = get_agent_display_info(agent_type)
                
                # 智能体标题
                agent_heading = doc.add_heading(f"{agent_info['icon']} {agent_info['name']}", level=2)
                agent_heading.runs[0].font.color.rgb = agent_info['color']
                
                # 处理结果数据
                if 'results' in agent_data and agent_data['results']:
                    for result in agent_data['results']:
                        content = extract_content_from_result(result)
                        if content:
                            formatted_content = format_content(content)
                            analysis_para = doc.add_paragraph(formatted_content)
                            analysis_para.style.font.size = Pt(11)
                
                # 添加状态信息
                if 'status' in agent_data:
                    status_para = doc.add_paragraph()
                    status_run = status_para.add_run(f"状态: {agent_data['status']}")
                    status_run.font.size = Pt(10)
                    status_run.font.color.rgb = RGBColor(149, 165, 166)
                    status_run.italic = True
                
                if 'start_time' in agent_data and 'end_time' in agent_data:
                    time_para = doc.add_paragraph()
                    time_run = time_para.add_run(f"执行时间: {agent_data['start_time']} - {agent_data['end_time']}")
                    time_run.font.size = Pt(10)
                    time_run.font.color.rgb = RGBColor(149, 165, 166)
                    time_run.italic = True
                
                doc.add_paragraph()  # 智能体之间的空行
        
        # 添加最终建议（如果有）
        if 'final_recommendation' in data and data['final_recommendation']:
            rec_heading = doc.add_heading('💡 最终投资建议', level=1)
            rec_heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
            
            rec_para = doc.add_paragraph(format_content(data['final_recommendation']))
            rec_para.style.font.size = Pt(12)
            
            # 高亮显示建议
            for run in rec_para.runs:
                if any(keyword in run.text for keyword in ['买入', '卖出', '持有', '建议', '推荐']):
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(231, 76, 60)
        
        # 添加风险提示
        risk_heading = doc.add_heading('⚠️ 风险提示', level=1)
        risk_heading.runs[0].font.color.rgb = RGBColor(231, 76, 60)
        
        risk_text = (
            "本报告仅供参考，不构成投资建议。投资有风险，入市需谨慎。\n"
            "请根据自身风险承受能力和投资目标做出投资决策。\n"
            "过往业绩不代表未来表现，市场价格可能大幅波动。"
        )
        risk_para = doc.add_paragraph(risk_text)
        risk_para.style.font.size = Pt(10)
        risk_para.style.font.color.rgb = RGBColor(231, 76, 60)
        
        # 添加页脚
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(149, 165, 166)
        footer_run.italic = True
        
        # 保存文档
        json_path = Path(json_file_path)
        output_path = json_path.parent / f"{json_path.stem}_fixed_report.docx"
        doc.save(output_path)
        
        print(f"✅ Word报告已生成: {output_path}")
        print(f"🎉 转换成功！可以用Microsoft Word打开查看: {output_path}")
        
        return str(output_path)
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    if len(sys.argv) != 2:
        print("使用方法: python json_to_word_converter.py <json_file_path>")
        print("示例: python json_to_word_converter.py progress_logs/session_20250729_161432.json")
        sys.exit(1)
    
    json_file_path = sys.argv[1]
    
    if not Path(json_file_path).exists():
        print(f"❌ 文件不存在: {json_file_path}")
        sys.exit(1)
    
    create_word_report(json_file_path)

if __name__ == "__main__":
    main()